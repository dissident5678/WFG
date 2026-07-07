#!/usr/bin/env python3
"""Build a subcontractor-facing bid packet from WFG opportunity artifacts.

This script is the deterministic side of the WFG subcontractor bid-packet
workflow. Hermes/subagents do the judgment-heavy document reading; this script
standardizes the packet, review files, hashes, and optional Google Drive review
bundle.

Safety boundary:
- It creates drafts and review artifacts only.
- It never sends email, contacts subcontractors, submits proposals, signs,
  certifies, spends money, or shares documents publicly.
- Missing/conflicting information goes to internal_review_summary.md, not the
  subcontractor-facing packet.
"""
from __future__ import annotations

import argparse
import copy
import datetime as dt
import hashlib
import json
import os
import re
import sys
from pathlib import Path
from typing import Any, Iterable

DEFAULT_OUTPUT_DIR = "subcontractor_bid_packet"
PROJECT = Path(os.environ.get("WFG_PROJECT_DIR", "/home/nick/workspace/wfg-gov-contracting-v2")).resolve()
TEMPLATE = Path(os.environ.get("WFG_SUB_BID_PACKET_TEMPLATE", str(PROJECT / "templates/subcontractor_bid_packet/WFG_Subcontractor_Bid_Packet_Template.docx"))).resolve()
INSTRUCTION_DOC = PROJECT / "templates/subcontractor_bid_packet/Hermes_Subcontractor_Bid_Packet_Instructions.docx"
TOKEN = Path(os.environ.get("GOOGLE_TOKEN_PATH", "/home/nick/.hermes/google_token.json"))
DRIVE_ROOT_FOLDER_ID = os.environ.get("WFG_DRIVE_ROOT_FOLDER_ID", "").strip()
SCOPES = [
    "https://www.googleapis.com/auth/drive.file",
]

INTERNAL_MARKERS = [
    "gate ", "approval", "not approved", "subcontractor not verified",
    "document missing", "user input required", "price not approved",
    "legal or compliance review required", "assumption — must be confirmed",
    "internal", "bid/no-bid", "margin", "markup", "profit strategy",
]

FIT_RULES = [
    {
        "name": "janitorial/custodial",
        "triggers": ["janitorial", "custodial", "cleaning", "floor care", "vct"],
        "required": ["commercial janitorial/custodial capability", "staffing and supervision for the stated frequency", "ability to support facility access and schedule requirements"],
        "preferred": ["Government/federal facility experience", "Office, lab, or secure-facility cleaning experience", "Wage determination familiarity if applicable"],
        "exclude": ["residential-only cleaning", "maid service only", "in-home estimates only"],
    },
    {
        "name": "parking lot striping/paving",
        "triggers": ["parking lot", "striping", "pavement marking", "asphalt", "sealcoat", "restriping"],
        "required": ["Commercial parking lot striping or pavement-marking capability", "Equipment and crew for the stated site size", "Ability to work around occupied facilities if required"],
        "preferred": ["Government/commercial site work", "Traffic-control capability", "Night/weekend work capability if required"],
        "exclude": ["Decorative painting only", "Residential driveway-only service", "Sign shop with no field striping crew"],
    },
    {
        "name": "electrical",
        "triggers": ["electrical", "electric", "lighting", "panel", "wiring", "breaker"],
        "required": ["Licensed commercial electrical contractor", "Service/repair capability matching the scope", "Ability to provide license and insurance information"],
        "preferred": ["Government/commercial facility work", "Prevailing wage familiarity where applicable", "Base/security access experience if required"],
        "exclude": ["Residential-only electrician", "Handyman-only service", "Low-voltage-only provider unless scope is low voltage"],
    },
    {
        "name": "solid waste / sludge hauling",
        "triggers": ["solid waste", "trash", "refuse", "dumpster", "recycling", "sludge", "wastewater", "biosolids", "lagoon"],
        "required": ["Commercial or municipal-scale waste/sludge hauling capability", "Transport/disposal capability matching the scope", "Permits/reporting support where required"],
        "preferred": ["Government/base facility service", "EPA/state/DOT permit familiarity", "Disposal documentation capability"],
        "exclude": ["Junk removal only", "Residential curbside-only provider", "Moving company only", "Septic-only provider without matching sludge capability"],
    },
    {
        "name": "landscaping/grounds/tree",
        "triggers": ["landscaping", "grounds", "mowing", "snow removal", "tree", "vegetation", "brush", "forestry", "mulch"],
        "required": ["Commercial grounds/tree/vegetation capability matching the scope", "Equipment and staffing for the stated acreage, quantity, or frequency", "Ability to support safety and access requirements"],
        "preferred": ["Government/commercial property work", "Traffic/property protection experience", "Debris disposal/chipping capability if required"],
        "exclude": ["Residential lawn care only", "One-person yard service without commercial capacity", "Tree trimming only if removal/clearing is required"],
    },
    {
        "name": "plumbing/mechanical",
        "triggers": ["plumbing", "plumber", "pipe", "backflow", "mechanical", "hvac", "boiler", "chiller"],
        "required": ["Licensed commercial plumbing/mechanical contractor", "Scope-specific service capability", "Ability to provide license and insurance information"],
        "preferred": ["Government/commercial facility work", "Emergency response if required", "Wage determination familiarity if applicable"],
        "exclude": ["Residential-only plumbing/HVAC", "Appliance repair only", "Unlicensed handyman service"],
    },
    {
        "name": "security/access control",
        "triggers": ["security", "guard", "access control", "surveillance", "camera", "alarm"],
        "required": ["Licensed commercial security/access-control provider matching the scope", "Credentialed staff if required", "Ability to support facility access requirements"],
        "preferred": ["Government/commercial facility work", "Cleared/credentialed staff if required", "Manufacturer certification where relevant"],
        "exclude": ["Residential alarm-only dealer", "Consumer camera installer only"],
    },
]


def utc_now() -> str:
    return dt.datetime.now(dt.timezone.utc).isoformat(timespec="seconds")


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def slug(value: str, max_len: int = 80) -> str:
    s = re.sub(r"[^a-zA-Z0-9._-]+", "-", str(value or "").strip().lower()).strip("-")
    return (s[:max_len].strip("-") or "item")


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore") if path.exists() else ""


def first_existing(opp: Path, *names: str) -> Path | None:
    candidates: list[Path] = []
    for name in names:
        candidates.extend([
            opp / name,
            opp / "drafts" / name,
            opp / "analysis" / name,
        ])
    for p in candidates:
        if p.exists():
            return p
    return None


def load_artifacts(opp: Path) -> dict[str, tuple[Path | None, str]]:
    mapping = {
        "brief": first_existing(opp, "02_SOLICITATION_BRIEF.md", "solicitation_brief.md"),
        "scope": first_existing(opp, "05_SCOPE_DECOMPOSITION.md", "scope_decomposition.md"),
        "criteria": first_existing(opp, "06_SUBCONTRACTOR_SOURCING_CRITERIA.md", "subcontractor_sourcing_criteria.md"),
        "missing": first_existing(opp, "04_MISSING_INFORMATION.md", "missing_information.md"),
        "manifest": first_existing(opp, "attachment_manifest.md"),
        "outreach": first_existing(opp, "07_DRAFT_OUTREACH.md", "07_DRAFT_SUBCONTRACTOR_QUOTE_REQUEST.md"),
    }
    return {k: (p, read_text(p) if p else "") for k, p in mapping.items()}


def section(text: str, heading: str) -> str:
    pattern = rf"^##\s+{re.escape(heading)}\s*$"
    m = re.search(pattern, text, flags=re.I | re.M)
    if not m:
        return ""
    start = m.end()
    nxt = re.search(r"^##\s+", text[start:], flags=re.M)
    end = start + nxt.start() if nxt else len(text)
    return text[start:end].strip()


def strip_md(line: str) -> str:
    line = re.sub(r"^[-*]\s+", "", line.strip())
    line = re.sub(r"^\d+\.\s+", "", line)
    line = re.sub(r"`([^`]+)`", r"\1", line)
    line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
    return line.strip()


def external_safe(value: str) -> bool:
    s = value.lower()
    if not value or "[" in value and "]" in value:
        return False
    return not any(marker in s for marker in INTERNAL_MARKERS)


def bullets(block: str, limit: int | None = None) -> list[str]:
    out: list[str] = []
    for raw in block.splitlines():
        s = strip_md(raw)
        if not s or s.startswith("#") or not external_safe(s):
            continue
        out.append(s)
    return out[:limit] if limit else out


def first_match(text: str, label: str, default: str = "") -> str:
    # Supports lines like "- Agency: ...", "Agency: ...", and simple tables.
    patterns = [
        rf"^-\s*\*\*{re.escape(label)}\*\*:\s*(.+)$",
        rf"^-\s*{re.escape(label)}:\s*(.+)$",
        rf"^{re.escape(label)}:\s*(.+)$",
        rf"^\|\s*{re.escape(label)}\s*\|\s*(.+?)\s*\|?$",
    ]
    for pat in patterns:
        m = re.search(pat, text, flags=re.I | re.M)
        if m:
            return strip_md(m.group(1).strip())
    m = re.search(rf"^##\s+{re.escape(label)}\s*$\n+([^#\n].+?)(?=\n##\s+|\Z)", text, flags=re.I | re.M | re.S)
    if m:
        val = " ".join(strip_md(x) for x in m.group(1).strip().splitlines() if strip_md(x))
        return val.strip()
    return default


def find_any(text: str, labels: Iterable[str], default: str = "") -> str:
    for label in labels:
        val = first_match(text, label, "")
        if val:
            return val
    return default


def parse_deadlineish(value: str) -> dt.datetime | None:
    if not value or not external_safe(value):
        return None
    raw = value.strip()
    # ISO-like values from SAM.gov.
    try:
        return dt.datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except Exception:
        pass
    # Common markdown text: 2026-07-22 14:00 ET / 07/22/2026 2:00 PM.
    candidates = [
        r"(20\d{2}-\d{1,2}-\d{1,2})(?:[ T]+(\d{1,2}:\d{2})(?:\s*(AM|PM))?)?",
        r"(\d{1,2}/\d{1,2}/20\d{2})(?:[ ,]+(\d{1,2}:\d{2})(?:\s*(AM|PM))?)?",
    ]
    for pat in candidates:
        m = re.search(pat, raw, flags=re.I)
        if not m:
            continue
        date_s, time_s, ampm = m.group(1), m.group(2), (m.group(3) or "")
        for date_fmt in ("%Y-%m-%d", "%m/%d/%Y"):
            try:
                d = dt.datetime.strptime(date_s, date_fmt).date()
                if time_s:
                    t = dt.datetime.strptime(time_s + (" " + ampm if ampm else ""), "%H:%M" + (" %p" if ampm else "")).time()
                else:
                    t = dt.time(17, 0)
                return dt.datetime.combine(d, t)
            except Exception:
                continue
    return None


def business_days_before(date_time: dt.datetime, days: int = 2) -> dt.datetime:
    d = date_time
    count = 0
    while count < days:
        d = d - dt.timedelta(days=1)
        if d.weekday() < 5:
            count += 1
    return d.replace(hour=12, minute=0, second=0, microsecond=0)


def derive_trade_filter(criteria: str, brief: str, scope: str) -> dict[str, Any]:
    blob = "\n".join([criteria, brief, scope]).lower()
    matches = [r for r in FIT_RULES if any(t in blob for t in r["triggers"])]
    matched = max(matches, key=lambda r: sum(1 for t in r["triggers"] if t in blob)) if matches else None
    if matched:
        return {
            "rule": matched["name"],
            "required_fit_terms": matched["required"],
            "preferred_experience": matched["preferred"],
            "exclude_or_manual_review": matched["exclude"] + ["Service area does not cover the place of performance", "Website/services do not show the required trade"],
        }
    trade = find_any(criteria, ["Trade/NAICS", "Trade", "NAICS"], "the specific trade/scope named in the opportunity")
    return {
        "rule": "generic trade-specific fit",
        "required_fit_terms": [f"Documented commercial capability for {trade}", "Service area covers the place of performance", "Scope-specific staffing/equipment capacity"],
        "preferred_experience": ["Similar government/commercial work", "Required licenses/insurance/permits for the scope", "Ability to meet schedule, access, wage, and reporting constraints"],
        "exclude_or_manual_review": ["Consumer-only or residential-only provider unless the solicitation is residential", "Provider offers a related but different service line", "Service area mismatch"],
    }


def detect(text: str, terms: Iterable[str]) -> bool:
    s = text.lower()
    return any(t.lower() in s for t in terms)


def list_from_manifest(manifest: str) -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    current: dict[str, str] = {}
    for raw in manifest.splitlines():
        s = raw.strip()
        if s.startswith("- Source URL:"):
            if current:
                out.append(current)
            current = {"source_url": s.split(":", 1)[1].strip()}
        elif "Filename:" in s:
            current["filename"] = s.split("Filename:", 1)[1].strip().strip("`")
        elif "Status:" in s:
            current["status"] = s.split("Status:", 1)[1].strip()
    if current:
        out.append(current)
    cleaned = []
    for idx, item in enumerate(out, start=1):
        name = item.get("filename") or item.get("source_url") or f"Attachment {idx}"
        if not name:
            continue
        purpose = "Reference for scope, pricing, site conditions, or compliance requirements."
        lower = name.lower()
        if any(x in lower for x in ["price", "pricing", "clin", "schedule"]):
            purpose = "Pricing format / line-item reference."
        elif any(x in lower for x in ["pws", "sow", "statement"]):
            purpose = "Scope of work reference."
        elif any(x in lower for x in ["wage", "wd"]):
            purpose = "Wage/labor requirement reference."
        elif any(x in lower for x in ["drawing", "plan", "spec"]):
            purpose = "Drawing/specification reference."
        cleaned.append({"attachment_name": name, "attachment_purpose": purpose, "review_required": "Yes"})
    return cleaned[:12]


def make_scope_items(brief: str, scope: str) -> list[dict[str, str]]:
    items = bullets(section(scope, "Work packages"), 30)
    if not items:
        items = bullets(section(brief, "Scope summary"), 30)
    if not items:
        items = bullets(section(scope, "Scope summary"), 30)
    if not items:
        one = first_match(brief, "Scope", "") or section(brief, "One-line summary") or "Price the trade scope described in the solicitation documents and attachments."
        items = [one]
    return [{"item_no": str(i + 1), "scope_description": x, "scope_source_reference": "Solicitation brief / scope decomposition"} for i, x in enumerate(items)]


def make_pricing_lines(brief: str, scope: str) -> list[dict[str, str]]:
    block = section(brief, "Price sheet") or section(brief, "CLINs") or section(scope, "Quote package needs") or ""
    lines = bullets(block, 30)
    pricing: list[dict[str, str]] = []
    for i, line in enumerate(lines, start=1):
        pricing.append({"line_id": str(i), "line_description": line, "unit": "", "quantity": "", "line_notes": "Price separately if applicable."})
    if not pricing:
        pricing.append({"line_id": "Base", "line_description": "Base scope shown in this packet", "unit": "LS", "quantity": "1", "line_notes": "Provide a lump-sum or solicitation-required breakout."})
    return pricing


def make_alternates(brief: str, scope: str) -> list[dict[str, str]]:
    block = section(brief, "Alternates") or section(brief, "Options") or section(scope, "Alternates") or section(scope, "Options")
    return [{"alternate_id": str(i + 1), "alternate_description": x, "alternate_unit_qty": "", "alternate_price_instruction": "Price separately"} for i, x in enumerate(bullets(block, 20))]


def make_exclusions(brief: str, scope: str) -> list[dict[str, str]]:
    block = section(scope, "Exclusions") or section(brief, "Exclusions") or ""
    return [{"excluded_item": x, "exclusion_note": "Do not include unless clarified in writing."} for x in bullets(block, 20)]


def make_clarifications(brief: str, scope: str, missing: str) -> list[dict[str, str]]:
    external_terms = ("site visit", "wage", "access", "credential", "hours", "schedule", "quantity", "sqft", "square", "frequency", "disposal", "drawing", "line item")
    found = [x for x in bullets(missing, 20) if any(t in x.lower() for t in external_terms)]
    found = [x for x in found if not any(t in x.lower() for t in ("proposal", "sam", "insurance status", "final technical", "manually verify"))]
    return [{"clarification_question": x, "clarification_reason": "Needed to tighten subcontractor pricing."} for x in found[:8]]


def build_packet_data(opp: Path, sub_quote_due_override: str = "") -> tuple[dict[str, Any], dict[str, Any], str]:
    artifacts = load_artifacts(opp)
    brief_path, brief = artifacts["brief"]
    scope_path, scope = artifacts["scope"]
    criteria_path, criteria = artifacts["criteria"]
    missing_path, missing = artifacts["missing"]
    manifest_path, manifest = artifacts["manifest"]
    all_text = "\n".join([brief, scope, criteria, missing, manifest])
    title = find_any(brief, ["Title", "Project title", "Opportunity / project"], "") or section(brief, "One-line summary") or opp.name
    agency = find_any(brief, ["Buyer/office", "Agency", "Agency / Owner", "Buyer"], "")
    solicitation = find_any(brief, ["Solicitation", "Solicitation number", "Solicitation / RFQ No.", "RFQ", "RFQ number"], "")
    notice = find_any(brief, ["Notice ID", "Notice"], "")
    if not notice:
        m = re.match(r"([0-9a-f]{12,32})[-_]", opp.name)
        notice = m.group(1) if m else ""
    government_due = find_any(brief, ["Quote due", "Response deadline", "Government due", "Government bid / quote due"], "")
    questions_due = find_any(brief, ["Questions due", "RFI deadline", "Question deadline"], "")
    site_visit = find_any(brief, ["Site visit", "Site visit date", "Pre-bid site visit"], "")
    place = find_any(brief, ["Place", "Place of performance", "Project location", "Location"], "")
    period = find_any(brief, ["POP", "Period of performance", "Performance period", "Expected period"], "")
    pricing_format = find_any(brief, ["Pricing format", "Requested pricing format", "Bid schedule", "Price sheet"], "") or "Use solicitation line items where shown; otherwise provide a clear trade breakout."
    gov_dt = parse_deadlineish(government_due)
    if sub_quote_due_override:
        sub_due = sub_quote_due_override
        sub_due_notes = "Subcontractor quote deadline set by WFG workflow."
    elif gov_dt:
        sub_due_dt = business_days_before(gov_dt, 2)
        sub_due = sub_due_dt.strftime("%Y-%m-%d %H:%M local project time")
        sub_due_notes = "Default internal deadline: two business days before the government due date."
    else:
        sub_due = "To be set by WFG before outreach"
        sub_due_notes = "Government due date could not be parsed from current extracted artifacts."
    filters = derive_trade_filter(criteria, brief, scope)
    scope_items = make_scope_items(brief, scope)
    pricing_lines = make_pricing_lines(brief, scope)
    alternates = make_alternates(brief, scope)
    exclusions = make_exclusions(brief, scope)
    clarifications = make_clarifications(brief, scope, missing)
    attachments = list_from_manifest(manifest)
    if not attachments:
        attachments = [{"attachment_name": "Solicitation package excerpts", "attachment_purpose": "Review source scope/pricing details if provided by WFG.", "review_required": "As needed"}]

    wage = detect(all_text, ["wage determination", "davis-bacon", "service contract act", "sca", "certified payroll", "labor classification"])
    bond_ins = detect(all_text, ["bond", "insurance", "certificate of insurance", "coi", "liability", "workers compensation"])
    access = detect(all_text, ["badge", "badging", "base access", "security", "escort", "credential", "background"])
    safety = detect(all_text, ["safety", "ppe", "em 385", "osha", "aha", "app", "site safety"])
    disposal = detect(all_text, ["dispose", "disposal", "haul", "dump", "debris", "recycling"])
    staging = detect(all_text, ["staging", "storage", "parking", "laydown"])
    utilities = detect(all_text, ["shutdown", "utility", "power outage", "water outage"])
    working_hours = find_any(brief, ["Working hours", "Work hours", "Hours"], "")
    scope_summary = "; ".join([x["scope_description"] for x in scope_items[:4]])
    if len(scope_summary) > 650:
        scope_summary = scope_summary[:647].rstrip() + "..."

    data: dict[str, Any] = {
        "prime_company_name": os.environ.get("WFG_PRIME_COMPANY_NAME", "Wright Foster Group LLC"),
        "prime_contact_name": os.environ.get("WFG_PRIME_CONTACT_NAME", "Nick Wright"),
        "prime_contact_email": os.environ.get("WFG_PRIME_CONTACT_EMAIL", "wrightfostergroup@gmail.com"),
        "prime_contact_phone": os.environ.get("WFG_PRIME_CONTACT_PHONE", "410-490-8681"),
        "project_title": title,
        "agency_name": agency or "Agency not listed in subcontractor packet; see solicitation reference.",
        "solicitation_number": solicitation or "Not listed",
        "notice_id": notice or "Not listed",
        "project_location": place or "See solicitation package",
        "trades_requested": filters["rule"],
        "sub_quote_due_datetime": sub_due,
        "sub_quote_due_notes": sub_due_notes,
        "government_due_datetime": government_due or "See solicitation package",
        "government_due_notes": "Final WFG response deadline. Subcontractor quotes are due earlier to allow review.",
        "questions_due_datetime": questions_due,
        "questions_due_notes": "Send subcontractor questions to WFG before this deadline so WFG can consolidate questions.",
        "questions_deadline_exists": bool(questions_due and external_safe(questions_due)),
        "site_visit_required": bool(site_visit and external_safe(site_visit) and "no " not in site_visit.lower()[:12]),
        "site_visit_datetime": site_visit,
        "site_visit_location": place or "project site",
        "site_visit_notes": site_visit,
        "award_timing_exists": bool(find_any(brief, ["Award timing", "Expected award", "NTP"], "")),
        "award_timing": find_any(brief, ["Award timing", "Expected award", "NTP"], ""),
        "award_timing_notes": "Estimated or stated by solicitation if shown.",
        "performance_period": period or "See solicitation package",
        "performance_notes": "Confirm availability and lead time in your quote.",
        "brief_scope_summary": scope_summary,
        "pricing_format": pricing_format,
        "price_valid_until_date": "30 days after quote unless your quote states otherwise",
        "scope_items": scope_items,
        "pricing_lines": pricing_lines,
        "alternates": alternates,
        "alternates_exist": bool(alternates),
        "exclusions": exclusions,
        "exclusions_exist": bool(exclusions),
        "attachments_for_subs": attachments,
        "clarification_questions": clarifications,
        "clarifications_requested_exist": bool(clarifications),
        "site_address_and_work_area": place or "See solicitation package",
        "access_security_relevant": access,
        "access_security_summary": "Confirm ability to meet site access, badging, escort, security, or credentialing requirements shown in the solicitation.",
        "working_hours_relevant": bool(working_hours),
        "working_hours": working_hours,
        "staging_relevant": staging,
        "staging_summary": "Include assumptions for parking, staging, storage, and laydown areas.",
        "utilities_relevant": utilities,
        "utilities_shutdown_summary": "Identify any utility shutdowns, outages, or owner coordination required for your work.",
        "disposal_relevant": disposal,
        "disposal_cleanup_summary": "Include debris, hauling, recycling, disposal, and final cleanup costs required by your scope.",
        "safety_relevant": safety,
        "safety_summary": "Include safety documentation, PPE, site safety, AHAs/APP/EM 385/OSHA support, and training assumptions if applicable.",
        "bond_or_insurance_relevant": bond_ins,
        "bond_insurance_summary": "Provide current insurance, bond, license, or certification details required for your trade and the solicitation.",
        "wage_determination_relevant": wage,
        "wage_determination_summary": "Price using the applicable wage/labor requirements shown in the solicitation, including wage determination or certified payroll requirements if applicable.",
        "fit_filter": filters,
        "generated_at": utc_now(),
        "opportunity_folder": str(opp),
        "template_path": str(TEMPLATE),
        "instruction_doc_path": str(INSTRUCTION_DOC),
    }

    missing_critical: list[str] = []
    for label, value in [
        ("agency_name", agency), ("solicitation_number", solicitation), ("notice_id", notice),
        ("project_location", place), ("government_due_datetime", government_due),
    ]:
        if not value or not external_safe(value):
            missing_critical.append(label)
    if sub_due.startswith("To be set"):
        missing_critical.append("sub_quote_due_datetime")
    if not scope_items:
        missing_critical.append("scope_items")

    conflicts: list[str] = []
    for term in ["conflict", "conflicting", "amendment", "changed", "superseded"]:
        if term in missing.lower():
            conflicts.append("Review missing-information file for possible amendment/conflict notes.")
            break

    internal_review = {
        "packet_readiness": "Needs human review" if missing_critical or conflicts else "Ready for approval review",
        "missing_critical_items": missing_critical,
        "conflicts_found": conflicts,
        "recommended_actions": [
            "Review subcontractor_bid_packet.docx before creating or approving outreach.",
            "Confirm the sub quote deadline and recipient list before Gate 2 approval.",
        ] + (["Manually verify missing critical fields before sending." ] if missing_critical else []),
        "source_files_used": {k: str(v[0]) for k, v in artifacts.items() if v[0]},
        "not_for_subcontractors": True,
    }

    source_map = {
        "project_title": {"source": str(brief_path) if brief_path else "opportunity folder", "method": "Title / one-line summary extraction"},
        "government_due_datetime": {"source": str(brief_path) if brief_path else "unknown", "method": "Quote due / response deadline extraction"},
        "questions_due_datetime": {"source": str(brief_path) if brief_path else "unknown", "method": "Questions due / RFI deadline extraction"},
        "scope_items": {"source": str(scope_path or brief_path or opp), "method": "Work packages / scope summary bullets"},
        "pricing_lines": {"source": str(brief_path or scope_path or opp), "method": "Price sheet / CLINs / quote package needs"},
        "attachments_for_subs": {"source": str(manifest_path) if manifest_path else "none", "method": "attachment_manifest.md parsing"},
        "missing_information": {"source": str(missing_path) if missing_path else "none", "method": "internal review only"},
    }
    return data, {"internal_review": internal_review, "source_map": source_map}, all_text


def truthy(data: dict[str, Any], key: str) -> bool:
    val = data.get(key)
    if isinstance(val, bool):
        return val
    if isinstance(val, (list, tuple, dict)):
        return bool(val)
    if val is None:
        return False
    s = str(val).strip().lower()
    return bool(s and s not in {"false", "0", "no", "none", "not listed", "not specified"})


def replace_fields(text: str, data: dict[str, Any], item: dict[str, Any] | None = None) -> str:
    ctx = dict(data)
    if item:
        ctx.update(item)
    def repl(m: re.Match[str]) -> str:
        key = m.group(1).strip()
        val = ctx.get(key, "")
        if isinstance(val, (list, dict)):
            return json.dumps(val, ensure_ascii=False)
        return str(val or "")
    return re.sub(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}", repl, text)


def render_conditionals(text: str, data: dict[str, Any]) -> str:
    # Run repeatedly so IF wrapping REPEAT rows cleans up before field replacement.
    prior = None
    out = text
    patterns = [
        (re.compile(r"\[\[IF_NOT\s+([A-Za-z0-9_]+)\]\]([\s\S]*?)\[\[/IF_NOT\]\]"), False),
        (re.compile(r"\[\[IF\s+([A-Za-z0-9_]+)\]\]([\s\S]*?)\[\[/IF\]\]"), True),
    ]
    while prior != out:
        prior = out
        for pattern, positive in patterns:
            def repl(m: re.Match[str]) -> str:
                cond = truthy(data, m.group(1))
                return m.group(2) if cond is positive else ""
            out = pattern.sub(repl, out)
    return out


def render_text(text: str, data: dict[str, Any], item: dict[str, Any] | None = None) -> str:
    out = render_conditionals(text, data)
    out = re.sub(r"\[\[/?REPEAT\s*[A-Za-z0-9_]*\]\]", "", out)
    out = replace_fields(out, data, item)
    out = re.sub(r"\s+\|\s+", " | ", out)
    return out.strip()


def set_cell_text(cell: Any, text: str) -> None:
    cell.text = text


def remove_paragraph(paragraph: Any) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def remove_row(row: Any) -> None:
    tr = row._tr
    tbl = tr.getparent()
    if tbl is not None:
        tbl.remove(tr)


def remove_table(table: Any) -> None:
    tbl = table._tbl
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)


def insert_row_after(row: Any) -> Any:
    new_tr = copy.deepcopy(row._tr)
    row._tr.addnext(new_tr)
    return new_tr


def render_docx_from_template(template: Path, out_docx: Path, data: dict[str, Any]) -> None:
    try:
        from docx import Document
        from docx.table import _Row
    except ModuleNotFoundError as exc:
        raise SystemExit("python-docx is required. Run in the Hermes environment with python-docx installed.") from exc
    if not template.exists():
        raise SystemExit(f"Subcontractor packet DOCX template not found: {template}")
    doc = Document(str(template))

    # Tables first because repeated rows may contain paragraph markers.
    for table in doc.tables:
        for row in list(table.rows):
            row_text = "\n".join(cell.text for cell in row.cells)
            rep = re.search(r"\[\[REPEAT\s+([A-Za-z0-9_]+)\]\]", row_text)
            if rep:
                list_name = rep.group(1)
                items = data.get(list_name) or []
                # Honor IF condition on the row as well.
                rendered_probe = render_conditionals(row_text, data)
                if not items or "[[REPEAT" not in rendered_probe and rep.group(0) in row_text and not truthy(data, list_name):
                    remove_row(row)
                    continue
                last_tr = row._tr
                for item in items:
                    new_tr = copy.deepcopy(row._tr)
                    last_tr.addnext(new_tr)
                    new_row = _Row(new_tr, table)
                    for cell in new_row.cells:
                        set_cell_text(cell, render_text(cell.text, data, item))
                    last_tr = new_tr
                remove_row(row)
                continue
            rendered_cells = [render_text(cell.text, data) for cell in row.cells]
            # Remove rows whose non-label cells became blank due to false IF sections.
            if any("[[IF " in cell.text or "[[IF_NOT" in cell.text for cell in row.cells):
                non_label = rendered_cells[1:] if len(rendered_cells) > 1 else rendered_cells
                if all(not x.strip() for x in non_label):
                    remove_row(row)
                    continue
            for cell, rendered in zip(row.cells, rendered_cells):
                set_cell_text(cell, rendered)

    def process_paragraphs(paragraphs):
        for paragraph in list(paragraphs):
            txt = paragraph.text
            if "[[" in txt or "{{" in txt:
                rendered = render_text(txt, data)
                if not rendered:
                    remove_paragraph(paragraph)
                else:
                    paragraph.text = rendered

    process_paragraphs(doc.paragraphs)

    # Remove optional sections entirely when the opportunity does not need them.
    optional_paragraph_rules = [
        ("3A. Alternates, Options, or Additive Items", "alternates_exist"),
        ("3B. Exclusions or Work Not Requested", "exclusions_exist"),
    ]
    for paragraph in list(doc.paragraphs):
        for prefix, cond in optional_paragraph_rules:
            if paragraph.text.strip().startswith(prefix) and not truthy(data, cond):
                remove_paragraph(paragraph)

    optional_table_rules = [
        ("Alt / Option", "alternates_exist"),
        ("Excluded Item", "exclusions_exist"),
        ("Question", "clarifications_requested_exist"),
    ]
    for table in list(doc.tables):
        header_text = " || ".join(cell.text for cell in table.rows[0].cells) if table.rows else ""
        for marker, cond in optional_table_rules:
            if marker in header_text and not truthy(data, cond):
                remove_table(table)
                break

    for section in doc.sections:
        process_paragraphs(section.header.paragraphs)
        process_paragraphs(section.footer.paragraphs)
        for part in [section.header, section.footer]:
            for table in part.tables:
                for row in list(table.rows):
                    rendered_cells = [render_text(cell.text, data) for cell in row.cells]
                    for cell, rendered in zip(row.cells, rendered_cells):
                        set_cell_text(cell, rendered)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def markdown_packet(data: dict[str, Any]) -> str:
    lines = [
        "# Subcontractor Bid Packet",
        "",
        str(data["project_title"]),
        "",
        "Prepared by Wright Foster Group LLC for subcontractor pricing review.",
        "",
        f"{data['prime_company_name']} is requesting subcontractor pricing for the work summarized below.",
        "",
        "## Quick Fit Check",
        "",
        "Before spending time on a quote, please confirm your company fits this work.",
        "",
        "Required service fit:",
    ]
    lines += [f"- {x}" for x in data.get("fit_filter", {}).get("required_fit_terms", [])]
    lines += ["", "Preferred experience:"]
    lines += [f"- {x}" for x in data.get("fit_filter", {}).get("preferred_experience", [])]
    lines += ["", "Not a fit / manual review before quoting:"]
    lines += [f"- {x}" for x in data.get("fit_filter", {}).get("exclude_or_manual_review", [])]
    lines += [
        "",
        "## Quote Request Snapshot",
        f"- Prime Contractor: {data['prime_company_name']}",
        f"- Prime Contact: {data['prime_contact_name']} | {data['prime_contact_email']} | {data['prime_contact_phone']}",
        f"- Agency / Owner: {data['agency_name']}",
        f"- Solicitation / RFQ No.: {data['solicitation_number']}",
        f"- Notice ID: {data['notice_id']}",
        f"- Project Location: {data['project_location']}",
        f"- Trades Requested: {data['trades_requested']}",
        f"- Sub Quote Due To Prime: {data['sub_quote_due_datetime']}",
        f"- Government Response Due: {data['government_due_datetime']}",
        "",
        "## Scope Summary",
        data["brief_scope_summary"],
        "",
        "## Work to Price",
    ]
    lines += [f"- {x['item_no']}. {x['scope_description']}" for x in data.get("scope_items", [])]
    lines += ["", "## Pricing Instructions", data["pricing_format"], ""]
    lines += ["| Line / CLIN | Description | Unit | Quantity | Notes |", "|---|---|---|---|---|"]
    for x in data.get("pricing_lines", []):
        lines.append(f"| {x.get('line_id','')} | {x.get('line_description','')} | {x.get('unit','')} | {x.get('quantity','')} | {x.get('line_notes','')} |")
    if data.get("site_visit_required") or data.get("questions_deadline_exists"):
        lines += ["", "## Key Dates"]
        if data.get("site_visit_required"):
            lines.append(f"- Site visit: {data.get('site_visit_datetime')} — {data.get('site_visit_notes')}")
        if data.get("questions_deadline_exists"):
            lines.append(f"- Questions due: {data.get('questions_due_datetime')}")
    lines += ["", "## Quote Submission Requirements", f"Return your quote to {data['prime_contact_name']} at {data['prime_contact_email']} by {data['sub_quote_due_datetime']}."]
    return "\n".join(lines).strip() + "\n"


def internal_review_markdown(review: dict[str, Any], data: dict[str, Any]) -> str:
    r = review["internal_review"]
    lines = [
        "# Internal Review Summary — Subcontractor Bid Packet",
        "",
        "Audience: Nick / WFG only. Do not send this file to subcontractors.",
        "",
        f"Packet readiness: {r['packet_readiness']}",
        f"Generated at: {data['generated_at']}",
        f"Opportunity folder: `{data['opportunity_folder']}`",
        "",
        "## Critical Missing Items",
    ]
    lines += [f"- {x}" for x in r["missing_critical_items"]] or ["- None identified from current extracted artifacts."]
    lines += ["", "## Conflicts Found"]
    lines += [f"- {x}" for x in r["conflicts_found"]] or ["- None identified from current extracted artifacts."]
    lines += ["", "## Recommended Actions"]
    lines += [f"- {x}" for x in r["recommended_actions"]]
    lines += ["", "## Source Files Used"]
    lines += [f"- {k}: `{v}`" for k, v in r["source_files_used"].items()] or ["- No extracted source files found."]
    lines += ["", "## Approval Gate Reminder", "- This packet is a draft. Do not send it or any outreach until Gate 2 external outreach approval is recorded for the exact recipient list, message, packet version, and packet hash."]
    return "\n".join(lines).strip() + "\n"


def google_drive_service():
    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaFileUpload
    except ModuleNotFoundError as exc:
        raise SystemExit("Google API libraries required for --drive. Install google-api-python-client and google-auth in the Hermes environment.") from exc
    creds = Credentials.from_authorized_user_file(str(TOKEN), scopes=SCOPES)
    if creds.expired and creds.refresh_token:
        creds.refresh(Request())
    return build("drive", "v3", credentials=creds, cache_discovery=False), MediaFileUpload


def drive_find_or_create_folder(drive: Any, name: str, parent_id: str | None = None) -> str:
    safe_name = name.replace("'", "\\'")
    q = f"mimeType='application/vnd.google-apps.folder' and name='{safe_name}' and trashed=false"
    if parent_id:
        q += f" and '{parent_id}' in parents"
    res = drive.files().list(q=q, fields="files(id,name)", spaces="drive", pageSize=10).execute()
    files = res.get("files", [])
    if files:
        return files[0]["id"]
    body: dict[str, Any] = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
    if parent_id:
        body["parents"] = [parent_id]
    created = drive.files().create(body=body, fields="id").execute()
    return created["id"]


def drive_upload_file(drive: Any, media_cls: Any, path: Path, folder_id: str) -> dict[str, Any]:
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if path.suffix.lower() == ".docx" else "text/markdown" if path.suffix.lower() == ".md" else "application/json"
    media = media_cls(str(path), mimetype=mime, resumable=False)
    existing = drive.files().list(q=f"name='{path.name.replace("'", "\\'")}' and '{folder_id}' in parents and trashed=false", fields="files(id,name)", spaces="drive", pageSize=10).execute().get("files", [])
    body = {"name": path.name, "parents": [folder_id]}
    if existing:
        fid = existing[0]["id"]
        meta = drive.files().update(fileId=fid, media_body=media, fields="id,name,webViewLink,mimeType").execute()
    else:
        meta = drive.files().create(body=body, media_body=media, fields="id,name,webViewLink,mimeType").execute()
    return meta


def create_drive_review_bundle(opp: Path, out_dir: Path, files: list[Path]) -> dict[str, Any]:
    try:
        import wfg_drive_review_hub
    except ModuleNotFoundError:
        sys.path.insert(0, str(Path(__file__).resolve().parent))
        import wfg_drive_review_hub

    return wfg_drive_review_hub.safe_upload_review_bundle(opp, extra_files=files)


def write_outputs(opp: Path, out_dir: Path, data: dict[str, Any], review: dict[str, Any], *, docx: bool, drive: bool, template: Path) -> dict[str, Any]:
    out_dir.mkdir(parents=True, exist_ok=True)
    md_path = out_dir / "subcontractor_bid_packet.md"
    internal_path = out_dir / "internal_review_summary.md"
    source_map_path = out_dir / "source_map.json"
    data_path = out_dir / "bid_packet_data.json"
    manifest_path = out_dir / "review_manifest.json"
    md_path.write_text(markdown_packet(data), encoding="utf-8")
    internal_path.write_text(internal_review_markdown(review, data), encoding="utf-8")
    source_map_path.write_text(json.dumps(review["source_map"], indent=2, sort_keys=True), encoding="utf-8")
    data_path.write_text(json.dumps(data, indent=2, sort_keys=True), encoding="utf-8")
    files = [md_path, internal_path, source_map_path, data_path]
    result: dict[str, Any] = {
        "markdown": str(md_path),
        "internal_review_summary": str(internal_path),
        "source_map": str(source_map_path),
        "bid_packet_data": str(data_path),
        "packet_readiness": review["internal_review"]["packet_readiness"],
        "missing_critical_items": review["internal_review"]["missing_critical_items"],
    }
    if docx:
        docx_path = out_dir / "subcontractor_bid_packet.docx"
        render_docx_from_template(template, docx_path, data)
        files.append(docx_path)
        result["docx"] = str(docx_path)
    content_hash = sha256_text("".join(p.read_text(errors="ignore") if p.suffix.lower() != ".docx" else hashlib.sha256(p.read_bytes()).hexdigest() for p in files))
    result["packet_hash"] = content_hash
    result["packet_version"] = "subpacket-" + content_hash[:16]
    if drive:
        drive_meta = create_drive_review_bundle(opp, out_dir, files)
        result["google_drive"] = drive_meta
        (out_dir / "google_drive_review_bundle.json").write_text(json.dumps(drive_meta, indent=2, sort_keys=True), encoding="utf-8")
    manifest = {
        "generated_at": utc_now(),
        "opportunity_folder": str(opp),
        "outputs": result,
        "human_approval_required_before_external_use": True,
        "safe_to_send_files_after_approval": [str(result.get("docx") or md_path)],
        "internal_only_files": [str(internal_path), str(source_map_path), str(data_path), str(manifest_path)],
    }
    manifest_path.write_text(json.dumps(manifest, indent=2, sort_keys=True), encoding="utf-8")
    result["review_manifest"] = str(manifest_path)
    return result


def build_packet(opp: Path) -> tuple[str, dict[str, Any]]:
    """Backward-compatible helper used by legacy tests."""
    data, review, _ = build_packet_data(opp)
    return markdown_packet(data), data


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("opportunity_folder", type=Path)
    ap.add_argument("--output-dir", default=DEFAULT_OUTPUT_DIR)
    ap.add_argument("--docx", action="store_true", help="Write subcontractor_bid_packet.docx from the dynamic DOCX template")
    ap.add_argument("--template", type=Path, default=TEMPLATE)
    ap.add_argument("--drive", action="store_true", help="Upload private review bundle to Google Drive review hub")
    ap.add_argument("--sub-quote-due", default="", help="Override subcontractor quote deadline shown in packet")
    ap.add_argument("--no-google-doc", action="store_true", help="Deprecated compatibility flag; use --drive for private Drive upload. Ignored.")
    ap.add_argument("--share-anyone", action="store_true", help="Deprecated unsafe flag. This script does not create public links.")
    args = ap.parse_args()
    opp = args.opportunity_folder.resolve()
    if not opp.exists():
        raise SystemExit(f"Opportunity folder not found: {opp}")
    if args.share_anyone:
        raise SystemExit("Refusing --share-anyone. WFG review bundles must stay private until explicit external approval and a separate send workflow.")
    data, review, _ = build_packet_data(opp, args.sub_quote_due)
    out_dir = opp / args.output_dir
    result = write_outputs(opp, out_dir, data, review, docx=args.docx, drive=args.drive, template=args.template.resolve())
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
